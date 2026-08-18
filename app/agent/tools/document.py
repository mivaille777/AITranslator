"""Local document tools used by the desktop Agent runtime."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from threading import RLock
from typing import Iterable

from app.agent.tools.base import ToolResult


SUPPORTED_DOCUMENT_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md", ".markdown"})
MAX_DOCUMENT_FILE_BYTES = 30 * 1024 * 1024
MAX_DOCUMENT_TEXT_CHARS = 500_000
DEFAULT_READ_CHARS = 12_000
DEFAULT_SUMMARY_CONTEXT_CHARS = 28_000
DEFAULT_SEARCH_RESULTS = 5


@dataclass(frozen=True, slots=True)
class DocumentChunk:
    index: int
    label: str
    text: str


@dataclass(frozen=True, slots=True)
class DocumentSession:
    path: Path
    kind: str
    text: str
    chunks: tuple[DocumentChunk, ...]

    @property
    def name(self) -> str:
        return self.path.name

    @property
    def character_count(self) -> int:
        return len(self.text)


class DocumentTools:
    """Open, extract and search one active local document.

    The active document lives only in process memory. Conversation persistence
    stays independent from local file contents so reopening AITranslator never
    silently re-reads a private file.
    """

    def __init__(self) -> None:
        self._lock = RLock()
        self._current: DocumentSession | None = None

    @property
    def current(self) -> DocumentSession | None:
        with self._lock:
            return self._current

    def open_file(self, path: str) -> ToolResult:
        resolved = Path(str(path)).expanduser()
        if not resolved.exists() or not resolved.is_file():
            return ToolResult("open_file", False, "文件不存在或不是普通文件。")
        suffix = resolved.suffix.lower()
        if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
            return ToolResult(
                "open_file",
                False,
                "暂不支持该文件类型。支持 PDF、DOCX、TXT 和 Markdown。",
            )
        try:
            size = resolved.stat().st_size
        except OSError as exc:
            return ToolResult("open_file", False, f"无法读取文件信息：{exc}")
        if size > MAX_DOCUMENT_FILE_BYTES:
            return ToolResult(
                "open_file",
                False,
                f"文件过大，当前上限为 {MAX_DOCUMENT_FILE_BYTES // (1024 * 1024)} MB。",
            )

        try:
            session = self._extract_session(resolved, suffix)
        except Exception as exc:
            return ToolResult(
                "open_file",
                False,
                f"文档解析失败：{type(exc).__name__}",
            )
        if not session.text.strip():
            return ToolResult("open_file", False, "文档中没有提取到可读文本。")
        with self._lock:
            self._current = session
        return ToolResult(
            "open_file",
            True,
            (
                f"已打开文档：{session.name}\n"
                f"类型：{session.kind}\n"
                f"提取字符数：{session.character_count}\n"
                f"可检索片段数：{len(session.chunks)}"
            ),
            {
                "document_name": session.name,
                "document_path": str(session.path),
                "document_kind": session.kind,
                "character_count": session.character_count,
                "chunk_count": len(session.chunks),
            },
        )

    def read_document(self, max_chars: int = DEFAULT_READ_CHARS, offset: int = 0) -> ToolResult:
        session = self.current
        if session is None:
            return ToolResult("read_document", False, "当前没有已打开的文档。")
        safe_max = min(50_000, max(500, int(max_chars)))
        safe_offset = max(0, int(offset))
        excerpt = session.text[safe_offset : safe_offset + safe_max]
        return ToolResult(
            "read_document",
            True,
            excerpt,
            {
                "document_name": session.name,
                "offset": safe_offset,
                "returned_chars": len(excerpt),
                "total_chars": session.character_count,
            },
        )

    def extract_document_text(self, max_chars: int = MAX_DOCUMENT_TEXT_CHARS) -> ToolResult:
        session = self.current
        if session is None:
            return ToolResult("extract_document_text", False, "当前没有已打开的文档。")
        safe_max = min(MAX_DOCUMENT_TEXT_CHARS, max(1_000, int(max_chars)))
        text = session.text[:safe_max]
        truncated = len(text) < session.character_count
        return ToolResult(
            "extract_document_text",
            True,
            text,
            {
                "document_name": session.name,
                "returned_chars": len(text),
                "total_chars": session.character_count,
                "truncated": truncated,
            },
        )

    def search_document(self, query: str, max_results: int = DEFAULT_SEARCH_RESULTS) -> ToolResult:
        session = self.current
        if session is None:
            return ToolResult("search_document", False, "当前没有已打开的文档。")
        normalized_query = " ".join(str(query or "").strip().split())
        if not normalized_query:
            return ToolResult("search_document", False, "文档检索关键词不能为空。")
        terms = [item.lower() for item in re.findall(r"[\w\-\.]+|[\u4e00-\u9fff]+", normalized_query)]
        if not terms:
            terms = [normalized_query.lower()]

        scored: list[tuple[int, DocumentChunk]] = []
        query_lower = normalized_query.lower()
        for chunk in session.chunks:
            haystack = chunk.text.lower()
            score = haystack.count(query_lower) * 12
            score += sum(haystack.count(term) for term in terms)
            if score > 0:
                scored.append((score, chunk))
        scored.sort(key=lambda item: (-item[0], item[1].index))
        selected = scored[: min(12, max(1, int(max_results)))]
        if not selected:
            return ToolResult(
                "search_document",
                True,
                f"在《{session.name}》中没有找到与“{normalized_query}”直接匹配的片段。",
                {"document_name": session.name, "query": normalized_query, "matches": 0},
            )

        sections: list[str] = []
        matches: list[dict[str, object]] = []
        for rank, (score, chunk) in enumerate(selected, 1):
            excerpt = self._search_excerpt(chunk.text, terms)
            sections.append(f"[{rank}] {chunk.label}\n{excerpt}")
            matches.append(
                {
                    "rank": rank,
                    "label": chunk.label,
                    "score": score,
                    "excerpt": excerpt,
                }
            )
        return ToolResult(
            "search_document",
            True,
            "\n\n".join(sections),
            {
                "document_name": session.name,
                "query": normalized_query,
                "matches": matches,
            },
        )

    def summarize_document(self, max_chars: int = DEFAULT_SUMMARY_CONTEXT_CHARS) -> ToolResult:
        """Return a bounded evidence package for LLM summarization.

        Summarization itself remains an LLM reasoning step; this deterministic
        tool owns file access, extraction and evidence bounding only.
        """

        session = self.current
        if session is None:
            return ToolResult("summarize_document", False, "当前没有已打开的文档。")
        safe_max = min(60_000, max(4_000, int(max_chars)))
        evidence = self._summary_evidence(session, safe_max)
        return ToolResult(
            "summarize_document",
            True,
            evidence,
            {
                "document_name": session.name,
                "document_kind": session.kind,
                "total_chars": session.character_count,
                "returned_chars": len(evidence),
                "requires_llm": True,
                "instruction": (
                    "请仅依据工具提供的文档证据，总结文档的主题、核心观点、方法/论证、"
                    "关键结论和重要限制；若证据被截断，要明确说明。"
                ),
            },
        )

    def _extract_session(self, path: Path, suffix: str) -> DocumentSession:
        if suffix == ".pdf":
            kind = "PDF"
            chunks = self._extract_pdf(path)
        elif suffix == ".docx":
            kind = "DOCX"
            chunks = self._extract_docx(path)
        else:
            kind = "Markdown" if suffix in {".md", ".markdown"} else "TXT"
            chunks = self._extract_text(path)
        cleaned = tuple(chunk for chunk in chunks if chunk.text.strip())
        joined = "\n\n".join(chunk.text.strip() for chunk in cleaned).strip()
        if len(joined) > MAX_DOCUMENT_TEXT_CHARS:
            joined = joined[:MAX_DOCUMENT_TEXT_CHARS]
        return DocumentSession(path=path.resolve(), kind=kind, text=joined, chunks=cleaned)

    @staticmethod
    def _extract_pdf(path: Path) -> tuple[DocumentChunk, ...]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks: list[DocumentChunk] = []
        for index, page in enumerate(reader.pages):
            text = str(page.extract_text() or "").strip()
            if text:
                chunks.append(DocumentChunk(index=index, label=f"PDF 第 {index + 1} 页", text=text))
        return tuple(chunks)

    @staticmethod
    def _extract_docx(path: Path) -> tuple[DocumentChunk, ...]:
        from docx import Document

        document = Document(str(path))
        chunks: list[DocumentChunk] = []
        index = 0
        for paragraph in document.paragraphs:
            text = str(paragraph.text or "").strip()
            if not text:
                continue
            style = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
            label = f"段落 {index + 1}"
            if style.lower().startswith("heading"):
                label = f"{style} · 段落 {index + 1}"
            chunks.append(DocumentChunk(index=index, label=label, text=text))
            index += 1
        for table_index, table in enumerate(document.tables, 1):
            rows: list[str] = []
            for row in table.rows:
                cells = [" ".join(str(cell.text or "").split()) for cell in row.cells]
                rows.append(" | ".join(cells))
            text = "\n".join(item for item in rows if item.strip()).strip()
            if text:
                chunks.append(DocumentChunk(index=index, label=f"表格 {table_index}", text=text))
                index += 1
        return tuple(chunks)

    @staticmethod
    def _extract_text(path: Path) -> tuple[DocumentChunk, ...]:
        raw: str | None = None
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                raw = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        if raw is None:
            raw = path.read_text(encoding="utf-8", errors="replace")
        paragraphs = [item.strip() for item in re.split(r"\n\s*\n", raw) if item.strip()]
        if not paragraphs:
            paragraphs = [raw.strip()] if raw.strip() else []
        return tuple(
            DocumentChunk(index=index, label=f"片段 {index + 1}", text=text)
            for index, text in enumerate(paragraphs)
        )

    @staticmethod
    def _search_excerpt(text: str, terms: Iterable[str], radius: int = 320) -> str:
        lowered = text.lower()
        positions = [lowered.find(term) for term in terms if term and lowered.find(term) >= 0]
        position = min(positions) if positions else 0
        start = max(0, position - radius)
        end = min(len(text), position + radius)
        excerpt = text[start:end].strip()
        if start > 0:
            excerpt = "…" + excerpt
        if end < len(text):
            excerpt += "…"
        return excerpt

    @staticmethod
    def _summary_evidence(session: DocumentSession, max_chars: int) -> str:
        if session.character_count <= max_chars:
            return session.text
        # Preserve both the beginning and end so conclusions are not dropped
        # from long papers while keeping the observation bounded for the LLM.
        head_size = max_chars * 2 // 3
        tail_size = max_chars - head_size
        return (
            session.text[:head_size].rstrip()
            + "\n\n[...文档中部已由工具截断...]\n\n"
            + session.text[-tail_size:].lstrip()
        )


__all__ = [
    "DEFAULT_READ_CHARS",
    "DEFAULT_SEARCH_RESULTS",
    "DEFAULT_SUMMARY_CONTEXT_CHARS",
    "DocumentChunk",
    "DocumentSession",
    "DocumentTools",
    "MAX_DOCUMENT_FILE_BYTES",
    "MAX_DOCUMENT_TEXT_CHARS",
    "SUPPORTED_DOCUMENT_EXTENSIONS",
]
