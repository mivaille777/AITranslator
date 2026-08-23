from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from backend.rag.exceptions import RagParsingError
from backend.rag.models import NormalizedDocument
from backend.rag.parsers.base import BaseFileParser, ParsedBlock, compose_blocks

_HEADING_STYLE = re.compile(r"^Heading\s+([1-6])$", re.IGNORECASE)


class DocxDocumentParser(BaseFileParser):
    name = "docx"
    version = "docx-v1"
    supported_suffixes = frozenset({".docx"})

    def parse(self, source: str | Path) -> NormalizedDocument:
        path = self._resolve_source(source)
        raw_bytes = self._read_bytes(path)
        try:
            docx = Document(str(path))
        except Exception as exc:
            raise RagParsingError(f"failed to parse DOCX document: {path}") from exc

        blocks: list[ParsedBlock] = []
        for paragraph in docx.paragraphs:
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            match = _HEADING_STYLE.match(style_name or "")
            heading_level = int(match.group(1)) if match else None
            blocks.append(ParsedBlock(text=text, heading_level=heading_level))

        text, sections = compose_blocks(blocks)
        if not text:
            raise RagParsingError(f"document contains no extractable text: {path}")

        properties = docx.core_properties
        title = (properties.title or "").strip() or path.stem
        metadata: dict[str, object] = {}
        if properties.author:
            metadata["author"] = str(properties.author)
        if properties.subject:
            metadata["subject"] = str(properties.subject)

        document = self._build_document(
            path=path,
            raw_bytes=raw_bytes,
            title=title,
            source_kind="docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata=metadata,
        )
        return NormalizedDocument(
            document=document,
            text=text,
            sections=sections,
            metadata={"parser_name": self.name, "parser_version": self.version},
        )


__all__ = ["DocxDocumentParser"]
