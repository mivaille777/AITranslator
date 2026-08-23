from pathlib import Path

from docx import Document

from backend.rag.parsers.docx import DocxDocumentParser


def test_docx_parser_preserves_heading_structure(tmp_path: Path) -> None:
    path = tmp_path / "paper.docx"
    document = Document()
    document.core_properties.title = "Controller Paper"
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Background.")
    document.add_heading("Method", level=2)
    document.add_paragraph("GP-guided PID tuning.")
    document.save(path)

    result = DocxDocumentParser().parse(path)

    assert result.document.title == "Controller Paper"
    assert result.document.source_kind == "docx"
    assert [section.heading for section in result.sections] == ["Introduction", "Method"]
    assert [section.level for section in result.sections] == [1, 2]
    assert "GP-guided PID tuning." in result.sections[1].text
