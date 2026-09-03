from __future__ import annotations

import base64
from pathlib import Path
from urllib.parse import unquote, urlparse
from urllib.request import url2pathname

from docx import Document
from docx.shared import Inches

from backend.rag.evidence_builder import build_evidence_item
from backend.rag.models import (
    DocumentElement,
    KnowledgeDocument,
    NormalizedDocument,
    RetrievalCandidate,
)
from backend.rag.multimodal import (
    MULTIMODAL_INDEX_VERSION,
    augment_document_with_visual_elements,
    build_multimodal_chunks,
    extract_visual_elements,
)

_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Zl1sAAAAASUVORK5CYII="
)


def _normalized(source: Path, *, source_kind: str) -> NormalizedDocument:
    return NormalizedDocument(
        document=KnowledgeDocument(
            document_id="doc-test",
            title="Demo paper",
            source_uri=source.as_uri(),
            source_kind=source_kind,
            mime_type="application/octet-stream",
            content_hash="abc123",
        ),
        text="Results\n\nFig. 1. Controller architecture.",
    )


def test_document_element_builds_retrievable_grounded_chunk(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF mock")
    visual = tmp_path / "figure.png"
    visual.write_bytes(_TINY_PNG)
    document = _normalized(source, source_kind="pdf").model_copy(
        update={
            "elements": [
                DocumentElement(
                    element_id="element-1",
                    document_id="doc-test",
                    modality="picture",
                    surrogate_text="Figure. Fig. 1. Controller architecture. Page: 2.",
                    page_number=2,
                    section_path=["Results"],
                    caption="Fig. 1. Controller architecture.",
                    asset_uri=visual.as_uri(),
                )
            ]
        }
    )

    chunks = build_multimodal_chunks(
        document,
        start_index=3,
        chunker_version=f"chunker+{MULTIMODAL_INDEX_VERSION}",
    )

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.chunk_type == "picture_element"
    assert chunk.text.startswith("Figure.")
    assert chunk.page_number == 2
    assert chunk.metadata["modality"] == "picture"
    assert chunk.metadata["asset_uri"] == visual.as_uri()
    assert chunk.metadata["retrieval_text_source"] == "surrogate_text"
    assert "base64" not in str(chunk.model_dump()).lower()

    evidence = build_evidence_item(RetrievalCandidate(chunk=chunk, rank=1))
    assert evidence.metadata["modality"] == "picture"
    assert evidence.metadata["asset_uri"] == visual.as_uri()
    assert evidence.metadata["visual_grounding_available"] is True


def test_docx_extraction_persists_picture_asset_and_caption(tmp_path: Path) -> None:
    source = tmp_path / "paper.docx"
    image = tmp_path / "source.png"
    image.write_bytes(_TINY_PNG)

    docx = Document()
    docx.add_paragraph("Results", style="Heading 1")
    docx.add_picture(str(image), width=Inches(0.2))
    docx.add_paragraph("Fig. 1. Controller architecture.")
    docx.save(source)

    normalized = _normalized(source, source_kind="docx")
    elements = extract_visual_elements(
        source,
        normalized,
        asset_root=tmp_path / "assets",
    )

    assert len(elements) == 1
    element = elements[0]
    assert element.modality == "picture"
    assert element.caption == "Fig. 1. Controller architecture."
    assert "Controller architecture" in element.surrogate_text
    assert element.asset_uri.startswith("file:")
    parsed_asset = urlparse(element.asset_uri)
    assert Path(url2pathname(unquote(parsed_asset.path))).exists()
    assert element.metadata["extraction_backend"] == "python-docx"
    assert element.metadata["layout_stable"] is False


def test_augmentation_keeps_text_contract_and_marks_visual_mode(tmp_path: Path) -> None:
    source = tmp_path / "paper.docx"
    image = tmp_path / "source.png"
    image.write_bytes(_TINY_PNG)
    docx = Document()
    docx.add_picture(str(image), width=Inches(0.2))
    docx.save(source)

    normalized = _normalized(source, source_kind="docx")
    augmented = augment_document_with_visual_elements(
        normalized,
        source,
        asset_root=tmp_path / "assets",
    )

    assert augmented.text == normalized.text
    assert len(augmented.elements) == 1
    assert augmented.metadata["multimodal_extraction_enabled"] is True
    assert augmented.metadata["visual_content_mode"] == "surrogate_text_and_asset"
    assert augmented.metadata["image_understanding_enabled"] is False
