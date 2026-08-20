"""Regression tests for local File / Document Agent tools."""

from __future__ import annotations

from docx import Document

from app.agent.tools.document import DocumentTools


def test_txt_document_can_open_read_search_and_prepare_summary(tmp_path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text(
        "Introduction\n\nLangGraph coordinates the agent workflow.\n\n"
        "Conclusion\n\nThe tool layer keeps deterministic file access outside the LLM.",
        encoding="utf-8",
    )
    tools = DocumentTools()

    opened = tools.open_file(str(path))
    assert opened.ok
    assert opened.metadata["document_kind"] == "TXT"

    read = tools.read_document()
    assert read.ok
    assert "LangGraph" in read.content

    searched = tools.search_document("deterministic file access")
    assert searched.ok
    assert searched.metadata["matches"]
    assert "deterministic file access" in searched.content

    summary = tools.summarize_document()
    assert summary.ok
    assert summary.metadata["requires_llm"] is True
    assert "Conclusion" in summary.content


def test_markdown_document_is_supported(tmp_path) -> None:
    path = tmp_path / "README.md"
    path.write_text("# Agent Tools\n\n- open_file\n- web_search\n", encoding="utf-8")
    tools = DocumentTools()

    result = tools.open_file(str(path))

    assert result.ok
    assert result.metadata["document_kind"] == "Markdown"
    assert tools.current is not None
    assert "web_search" in tools.current.text


def test_docx_document_extracts_paragraphs_and_tables(tmp_path) -> None:
    path = tmp_path / "paper.docx"
    document = Document()
    document.add_heading("Method", level=1)
    document.add_paragraph("The proposed method uses a bounded Agent tool runtime.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    document.save(path)

    tools = DocumentTools()
    opened = tools.open_file(str(path))
    extracted = tools.extract_document_text()

    assert opened.ok
    assert opened.metadata["document_kind"] == "DOCX"
    assert extracted.ok
    assert "Method" in extracted.content
    assert "Metric | Value" in extracted.content


def test_document_tools_reject_unsupported_extension(tmp_path) -> None:
    path = tmp_path / "data.csv"
    path.write_text("a,b", encoding="utf-8")

    result = DocumentTools().open_file(str(path))

    assert not result.ok
    assert "PDF" in result.content
